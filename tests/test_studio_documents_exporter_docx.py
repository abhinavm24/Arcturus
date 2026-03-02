"""Tests for core/studio/documents/exporter_docx.py — DOCX rendering."""

import pytest
from pathlib import Path

from docx import Document
from docx.shared import Pt

from core.schemas.studio_schema import DocumentContentTree, DocumentSection
from core.studio.documents.exporter_docx import export_to_docx


@pytest.fixture
def sample_content_tree():
    return DocumentContentTree(
        doc_title="Test Report",
        doc_type="report",
        abstract="This is a test abstract for the report.",
        sections=[
            DocumentSection(
                id="sec1",
                heading="Introduction",
                level=1,
                content="Introduction paragraph one.\n\nIntroduction paragraph two.",
                subsections=[
                    DocumentSection(
                        id="sec1a",
                        heading="Background",
                        level=2,
                        content="Background details with [ref1] citation.",
                    )
                ],
                citations=["ref1"],
            ),
            DocumentSection(
                id="sec2",
                heading="Findings",
                level=1,
                content="Key findings are presented here.",
            ),
            DocumentSection(
                id="sec3",
                heading="Conclusion",
                level=1,
                content="Final conclusions.",
            ),
        ],
        bibliography=[
            {"key": "ref1", "title": "Source A", "author": "Author A", "year": "2024"},
            {"key": "ref2", "title": "Source B", "author": "Author B", "url": "https://example.com"},
        ],
    )


@pytest.fixture
def output_path(tmp_path):
    return tmp_path / "test_output.docx"


class TestExportToDocx:
    def test_creates_file(self, sample_content_tree, output_path):
        result = export_to_docx(sample_content_tree, output_path)
        assert result == output_path
        assert output_path.exists()
        assert output_path.stat().st_size > 0

    def test_creates_parent_directories(self, sample_content_tree, tmp_path):
        nested_path = tmp_path / "sub" / "dir" / "test.docx"
        export_to_docx(sample_content_tree, nested_path)
        assert nested_path.exists()

    def test_contains_title(self, sample_content_tree, output_path):
        export_to_docx(sample_content_tree, output_path)
        doc = Document(str(output_path))
        texts = [p.text for p in doc.paragraphs]
        assert any("Test Report" in t for t in texts)

    def test_contains_abstract(self, sample_content_tree, output_path):
        export_to_docx(sample_content_tree, output_path)
        doc = Document(str(output_path))
        texts = [p.text for p in doc.paragraphs]
        assert any("test abstract" in t for t in texts)

    def test_contains_headings(self, sample_content_tree, output_path):
        export_to_docx(sample_content_tree, output_path)
        doc = Document(str(output_path))
        headings = [
            p.text for p in doc.paragraphs
            if p.style and p.style.name and p.style.name.startswith("Heading")
        ]
        assert "Introduction" in headings
        assert "Background" in headings
        assert "Findings" in headings
        assert "Bibliography" in headings

    def test_contains_section_content(self, sample_content_tree, output_path):
        export_to_docx(sample_content_tree, output_path)
        doc = Document(str(output_path))
        texts = [p.text for p in doc.paragraphs]
        all_text = " ".join(texts)
        assert "Introduction paragraph one" in all_text
        assert "Introduction paragraph two" in all_text

    def test_contains_bibliography_entries(self, sample_content_tree, output_path):
        export_to_docx(sample_content_tree, output_path)
        doc = Document(str(output_path))
        all_text = " ".join(p.text for p in doc.paragraphs)
        assert "[ref1]" in all_text
        assert "Author A" in all_text
        assert "Source A" in all_text

    def test_sanitizes_control_characters(self, tmp_path):
        tree = DocumentContentTree(
            doc_title="Control\x00Char Test",
            doc_type="report",
            sections=[
                DocumentSection(
                    id="s1", heading="Test\x0bHeading", level=1,
                    content="Content with\x08bad chars.",
                )
            ],
        )
        path = tmp_path / "sanitized.docx"
        export_to_docx(tree, path)
        doc = Document(str(path))
        all_text = " ".join(p.text for p in doc.paragraphs)
        assert "\x00" not in all_text
        assert "\x0b" not in all_text

    def test_empty_bibliography(self, tmp_path):
        tree = DocumentContentTree(
            doc_title="No Bib",
            doc_type="report",
            sections=[
                DocumentSection(id="s1", heading="Intro", level=1, content="Content."),
            ],
            bibliography=[],
        )
        path = tmp_path / "no_bib.docx"
        export_to_docx(tree, path)
        doc = Document(str(path))
        headings = [
            p.text for p in doc.paragraphs
            if p.style and p.style.name and p.style.name.startswith("Heading")
        ]
        assert "Bibliography" not in headings

    def test_no_abstract(self, tmp_path):
        tree = DocumentContentTree(
            doc_title="No Abstract",
            doc_type="report",
            abstract=None,
            sections=[
                DocumentSection(id="s1", heading="Intro", level=1, content="Content."),
            ],
        )
        path = tmp_path / "no_abstract.docx"
        export_to_docx(tree, path)
        assert path.exists()

    def test_docx_renders_bold_text(self, tmp_path):
        """Bold markdown should produce bold runs without raw ** markers."""
        tree = DocumentContentTree(
            doc_title="Bold Test",
            doc_type="report",
            sections=[
                DocumentSection(
                    id="s1", heading="Section", level=1,
                    content="This has **bold text** in it.",
                ),
            ],
        )
        path = tmp_path / "bold.docx"
        export_to_docx(tree, path)
        doc = Document(str(path))
        all_text = " ".join(p.text for p in doc.paragraphs)
        assert "bold text" in all_text
        assert "**bold text**" not in all_text
        # Check that at least one run is bold
        bold_found = False
        for p in doc.paragraphs:
            for run in p.runs:
                if run.bold and "bold" in (run.text or ""):
                    bold_found = True
        assert bold_found, "Expected a bold run containing 'bold'"

    def test_docx_renders_code_block(self, tmp_path):
        """Fenced code blocks should render without ``` markers."""
        tree = DocumentContentTree(
            doc_title="Code Test",
            doc_type="report",
            sections=[
                DocumentSection(
                    id="s1", heading="Section", level=1,
                    content="Example:\n\n```python\nprint('hello')\n```",
                ),
            ],
        )
        path = tmp_path / "code.docx"
        export_to_docx(tree, path)
        doc = Document(str(path))
        all_text = " ".join(p.text for p in doc.paragraphs)
        assert "print" in all_text
        assert "```" not in all_text

    def test_docx_renders_bullet_list(self, tmp_path):
        """Bullet list items should be present in the output."""
        tree = DocumentContentTree(
            doc_title="List Test",
            doc_type="report",
            sections=[
                DocumentSection(
                    id="s1", heading="Section", level=1,
                    content="Items:\n\n- Alpha\n- Beta\n- Gamma",
                ),
            ],
        )
        path = tmp_path / "list.docx"
        export_to_docx(tree, path)
        doc = Document(str(path))
        all_text = " ".join(p.text for p in doc.paragraphs)
        assert "Alpha" in all_text
        assert "Beta" in all_text
        assert "Gamma" in all_text

    def test_deeply_nested_sections(self, tmp_path):
        tree = DocumentContentTree(
            doc_title="Nested",
            doc_type="report",
            sections=[
                DocumentSection(
                    id="s1", heading="Level 1", level=1,
                    content="L1 content.",
                    subsections=[
                        DocumentSection(
                            id="s1a", heading="Level 2", level=2,
                            content="L2 content.",
                            subsections=[
                                DocumentSection(
                                    id="s1a1", heading="Level 3", level=3,
                                    content="L3 content.",
                                )
                            ],
                        )
                    ],
                )
            ],
        )
        path = tmp_path / "nested.docx"
        export_to_docx(tree, path)
        doc = Document(str(path))
        headings = [
            p.text for p in doc.paragraphs
            if p.style and p.style.name and p.style.name.startswith("Heading")
        ]
        assert "Level 1" in headings
        assert "Level 2" in headings
        assert "Level 3" in headings

    def test_docx_mermaid_shows_diagram_label(self, tmp_path):
        """Mermaid blocks should show 'Diagram (source)' label, not raw fences."""
        tree = DocumentContentTree(
            doc_title="Mermaid Test",
            doc_type="report",
            sections=[
                DocumentSection(
                    id="s1", heading="Architecture", level=1,
                    content="Overview:\n\n```mermaid\ngraph TD\n  A-->B\n```",
                ),
            ],
        )
        path = tmp_path / "mermaid.docx"
        export_to_docx(tree, path)
        doc = Document(str(path))
        all_text = " ".join(p.text for p in doc.paragraphs)
        assert "Diagram (source)" in all_text
        assert "mermaid.live" in all_text
        assert "graph TD" in all_text
        assert "```" not in all_text

    def test_docx_mermaid_has_styled_runs(self, tmp_path):
        """Mermaid header should be bold, code should be monospace."""
        tree = DocumentContentTree(
            doc_title="Mermaid Styles",
            doc_type="report",
            sections=[
                DocumentSection(
                    id="s1", heading="Arch", level=1,
                    content="```mermaid\ngraph TD\n  A-->B\n```",
                ),
            ],
        )
        path = tmp_path / "mermaid_styled.docx"
        export_to_docx(tree, path)
        doc = Document(str(path))
        # Find the header run
        header_found = False
        mono_found = False
        for p in doc.paragraphs:
            for run in p.runs:
                if "Diagram (source)" in (run.text or ""):
                    assert run.bold
                    header_found = True
                if "graph TD" in (run.text or ""):
                    assert run.font.name == "Courier New"
                    mono_found = True
        assert header_found, "Expected a bold 'Diagram (source)' run"
        assert mono_found, "Expected a monospace run with mermaid source"

    def test_loose_list_no_empty_paragraphs(self, tmp_path):
        """Loose lists (items separated by blank lines) should not create empty paragraphs."""
        tree = DocumentContentTree(
            doc_title="Loose List Test",
            doc_type="report",
            sections=[
                DocumentSection(
                    id="s1", heading="Section", level=1,
                    # Loose list: blank lines between items produce <li><p>text</p></li>
                    content="- Alpha\n\n- Beta\n\n- Gamma",
                ),
            ],
        )
        path = tmp_path / "loose_list.docx"
        export_to_docx(tree, path)
        doc = Document(str(path))

        # Collect list-styled paragraphs
        bullet_paras = [
            p for p in doc.paragraphs
            if p.style and p.style.name == "List Bullet"
        ]
        assert len(bullet_paras) == 3, f"Expected 3 bullet paragraphs, got {len(bullet_paras)}"

        # All bullet paragraphs should have content (no empty orphans)
        for bp in bullet_paras:
            assert bp.text.strip(), f"Found empty List Bullet paragraph"

        # Verify the content is correct
        bullet_texts = [p.text.strip() for p in bullet_paras]
        assert bullet_texts == ["Alpha", "Beta", "Gamma"]

        # No Normal-styled empty paragraphs should exist between list items
        empty_normal = [
            p for p in doc.paragraphs
            if p.style and p.style.name == "Normal" and not p.text.strip()
        ]
        assert len(empty_normal) == 0, (
            f"Found {len(empty_normal)} empty Normal paragraphs (orphans from loose list)"
        )

    def test_abstract_no_spacer_paragraphs(self, tmp_path):
        """Abstract should use paragraph spacing, not empty spacer paragraphs."""
        tree = DocumentContentTree(
            doc_title="Spacing Test",
            doc_type="report",
            abstract="This is the abstract.",
            sections=[
                DocumentSection(id="s1", heading="Intro", level=1, content="Content."),
            ],
        )
        path = tmp_path / "spacing.docx"
        export_to_docx(tree, path)
        doc = Document(str(path))

        # Find the abstract paragraph
        abstract_para = None
        for p in doc.paragraphs:
            if "This is the abstract." in p.text:
                abstract_para = p
                break
        assert abstract_para is not None, "Abstract paragraph not found"

        # Verify spacing attributes are set (not empty spacer paragraphs)
        assert abstract_para.paragraph_format.space_before == Pt(12), (
            f"Expected space_before=Pt(12), got {abstract_para.paragraph_format.space_before}"
        )
        assert abstract_para.paragraph_format.space_after == Pt(12), (
            f"Expected space_after=Pt(12), got {abstract_para.paragraph_format.space_after}"
        )

        # No empty spacer paragraphs adjacent to abstract
        all_texts = [p.text for p in doc.paragraphs]
        abstract_idx = next(i for i, t in enumerate(all_texts) if "This is the abstract." in t)
        # Paragraph before abstract should be the title, not empty
        if abstract_idx > 0:
            assert all_texts[abstract_idx - 1].strip(), (
                "Found empty spacer paragraph before abstract"
            )
