"""DOCX exporter — render DocumentContentTree to a Word document."""

from __future__ import annotations

import re
from html.parser import HTMLParser
from pathlib import Path
from typing import List

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

from core.schemas.studio_schema import DocumentContentTree, DocumentSection
from core.studio.documents.markdown_render import markdown_to_html


# XML-invalid control characters (except tab, newline, carriage return)
_CTRL_CHAR_RE = re.compile(r"[\x00-\x08\x0b\x0c\x0e-\x1f\x7f]")


def _sanitize(text: str) -> str:
    """Strip XML-invalid control characters."""
    return _CTRL_CHAR_RE.sub("", text)


class _DocxHtmlRenderer(HTMLParser):
    """Parse markdown-generated HTML and emit python-docx elements.

    Handles: <p>, <strong>/<b>, <em>/<i>, <code>, <pre>,
    <ul>/<ol> + <li>, <table>/<tr>/<th>/<td>, <blockquote>.
    """

    def __init__(self, doc: Document) -> None:
        super().__init__()
        self._doc = doc
        self._para = None  # current paragraph
        self._tag_stack: list[str] = []
        # Formatting state
        self._bold = False
        self._italic = False
        self._code = False
        self._pre = False
        # List state
        self._list_stack: list[str] = []  # "ul" or "ol"
        self._ol_counter: list[int] = []
        # Table state
        self._in_table = False
        self._table_rows: list[list[str]] = []
        self._current_row: list[str] = []
        self._current_cell_text = ""
        self._is_header_row = False
        # Blockquote
        self._in_blockquote = False
        # Loose-list state: <li><p>text</p></li> should reuse the <li> paragraph
        self._in_li = False
        # Mermaid diagram source blocks
        self._in_mermaid = False
        self._mermaid_zone = ""  # "header", "code", "footer", or ""

    def _ensure_paragraph(self) -> None:
        if self._para is None:
            self._para = self._doc.add_paragraph()

    def _add_run(self, text: str) -> None:
        text = _sanitize(text)
        if not text:
            return
        self._ensure_paragraph()
        run = self._para.add_run(text)
        if self._bold:
            run.bold = True
        if self._italic or self._in_blockquote:
            run.italic = True
        if self._code or self._pre:
            run.font.name = "Courier New"
            run.font.size = Pt(9)

    @staticmethod
    def _get_attr(attrs: list, name: str) -> str:
        """Get an attribute value from the tag attrs list."""
        for k, v in attrs:
            if k == name:
                return v or ""
        return ""

    def handle_starttag(self, tag: str, attrs: list) -> None:
        tag = tag.lower()
        self._tag_stack.append(tag)
        cls = self._get_attr(attrs, "class")

        # Mermaid diagram source container
        if tag == "div" and cls == "mermaid-source":
            self._in_mermaid = True
            return
        if self._in_mermaid and tag == "div":
            if cls == "mermaid-header":
                self._mermaid_zone = "header"
                self._para = self._doc.add_paragraph()
                self._para.paragraph_format.left_indent = Inches(0.3)
            elif cls == "mermaid-footer":
                self._mermaid_zone = "footer"
                self._para = self._doc.add_paragraph()
                self._para.paragraph_format.left_indent = Inches(0.3)
            return
        if self._in_mermaid and tag == "pre" and "mermaid-code" in cls:
            self._mermaid_zone = "code"
            self._pre = True
            self._para = self._doc.add_paragraph()
            self._para.paragraph_format.left_indent = Inches(0.3)
            return
        if self._in_mermaid and tag in ("code", "a"):
            return

        if tag == "p":
            if self._in_table:
                return
            if self._in_li and self._para is not None:
                return  # Reuse <li> paragraph for loose-list <p>
            self._para = self._doc.add_paragraph()
            if self._in_blockquote:
                self._para.paragraph_format.left_indent = Inches(0.5)
        elif tag in ("strong", "b"):
            self._bold = True
        elif tag in ("em", "i"):
            self._italic = True
        elif tag == "code":
            self._code = True
        elif tag == "pre":
            self._pre = True
            self._para = self._doc.add_paragraph()
        elif tag == "ul":
            self._list_stack.append("ul")
        elif tag == "ol":
            self._list_stack.append("ol")
            self._ol_counter.append(0)
        elif tag == "li":
            self._in_li = True
            self._para = self._doc.add_paragraph()
            if self._list_stack:
                list_type = self._list_stack[-1]
                if list_type == "ul":
                    self._para.style = self._doc.styles["List Bullet"]
                else:
                    self._para.style = self._doc.styles["List Number"]
                    self._ol_counter[-1] += 1
        elif tag == "table":
            self._in_table = True
            self._table_rows = []
        elif tag == "tr":
            self._current_row = []
            self._is_header_row = False
        elif tag in ("td", "th"):
            self._current_cell_text = ""
            if tag == "th":
                self._is_header_row = True
        elif tag == "blockquote":
            self._in_blockquote = True

    def handle_endtag(self, tag: str) -> None:
        tag = tag.lower()
        if self._tag_stack and self._tag_stack[-1] == tag:
            self._tag_stack.pop()

        # Mermaid container end
        if self._in_mermaid:
            if tag == "div" and self._mermaid_zone in ("header", "footer"):
                self._para = None
                self._mermaid_zone = ""
                return
            if tag == "pre" and self._mermaid_zone == "code":
                self._pre = False
                self._para = None
                self._mermaid_zone = ""
                return
            if tag == "div" and not self._mermaid_zone:
                # Closing the outer mermaid-source div
                self._in_mermaid = False
                return
            if tag in ("code", "a"):
                return

        if tag == "p":
            if not self._in_table and not self._in_li:
                self._para = None
        elif tag in ("strong", "b"):
            self._bold = False
        elif tag in ("em", "i"):
            self._italic = False
        elif tag == "code":
            self._code = False
        elif tag == "pre":
            self._pre = False
            self._para = None
        elif tag == "ul":
            if self._list_stack:
                self._list_stack.pop()
        elif tag == "ol":
            if self._list_stack:
                self._list_stack.pop()
            if self._ol_counter:
                self._ol_counter.pop()
        elif tag == "li":
            self._in_li = False
            self._para = None
        elif tag in ("td", "th"):
            self._current_row.append(self._current_cell_text.strip())
        elif tag == "tr":
            self._table_rows.append(self._current_row)
        elif tag == "table":
            self._flush_table()
            self._in_table = False
        elif tag == "blockquote":
            self._in_blockquote = False

    def handle_data(self, data: str) -> None:
        if self._in_table and ("td" in self._tag_stack or "th" in self._tag_stack):
            self._current_cell_text += data
            return
        # Skip whitespace-only text nodes inside list containers (inter-tag HTML formatting)
        if self._list_stack and not data.strip():
            return
        # Mermaid zones get distinct formatting
        if self._in_mermaid:
            text = _sanitize(data)
            if not text:
                return
            if self._mermaid_zone == "header":
                self._ensure_paragraph()
                run = self._para.add_run(text)
                run.bold = True
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(0x33, 0x66, 0x99)
                return
            if self._mermaid_zone == "footer":
                self._ensure_paragraph()
                run = self._para.add_run(text)
                run.font.size = Pt(8)
                run.font.color.rgb = RGBColor(0x66, 0x88, 0xAA)
                return
            if self._mermaid_zone == "code":
                self._ensure_paragraph()
                run = self._para.add_run(text)
                run.font.name = "Courier New"
                run.font.size = Pt(8)
                run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
                return
            return
        self._add_run(data)

    def _flush_table(self) -> None:
        """Emit accumulated table rows as a python-docx table."""
        if not self._table_rows:
            return
        num_cols = max(len(row) for row in self._table_rows)
        if num_cols == 0:
            return
        table = self._doc.add_table(
            rows=len(self._table_rows), cols=num_cols
        )
        table.style = "Table Grid"
        for r_idx, row_data in enumerate(self._table_rows):
            for c_idx, cell_text in enumerate(row_data):
                if c_idx < num_cols:
                    cell = table.cell(r_idx, c_idx)
                    cell.text = _sanitize(cell_text)
                    if r_idx == 0 and self._table_rows[0]:
                        # Bold header row
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                run.bold = True
        self._table_rows = []
        self._para = None


_EMPTY_P_RE = re.compile(r"<p>\s*</p>")


def _render_markdown_content(doc: Document, text: str) -> None:
    """Convert markdown text to HTML, then render into the DOCX document."""
    html = markdown_to_html(text)
    if not html:
        return
    html = _EMPTY_P_RE.sub("", html)
    renderer = _DocxHtmlRenderer(doc)
    renderer.feed(html)


def export_to_docx(
    content_tree: DocumentContentTree,
    output_path: Path,
) -> Path:
    """Render a DocumentContentTree to a DOCX file.

    Uses python-docx with default Word styles.
    Returns the output path.
    """
    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()

    # Set default font size
    style = doc.styles["Normal"]
    font = style.font
    font.size = Pt(11)

    # Title
    title_para = doc.add_heading(_sanitize(content_tree.doc_title), level=0)
    title_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Abstract
    if content_tree.abstract:
        abstract_para = doc.add_paragraph()
        abstract_para.paragraph_format.space_before = Pt(12)
        abstract_para.paragraph_format.space_after = Pt(12)
        abstract_para.style = doc.styles["Normal"]
        run = abstract_para.add_run(_sanitize(content_tree.abstract))
        run.italic = True

    # Sections
    _render_sections(doc, content_tree.sections)

    # Bibliography
    if content_tree.bibliography:
        doc.add_heading("Bibliography", level=1)
        for entry in content_tree.bibliography:
            key = entry.get("key", "")
            title = entry.get("title", "")
            author = entry.get("author", "")
            bib_text = f"[{key}] {author}. {title}."
            # Add extra fields if present
            year = entry.get("year", "")
            url = entry.get("url", "")
            if year:
                bib_text += f" ({year})"
            if url:
                bib_text += f" {url}"
            doc.add_paragraph(_sanitize(bib_text))

    doc.save(str(output_path))
    return output_path


def _render_sections(doc: Document, sections: List[DocumentSection]) -> None:
    """Recursively render sections as headings + rich content."""
    for section in sections:
        # Heading level: 1-3 (Word supports up to heading 9, but we cap at 3)
        heading_level = min(section.level, 3)
        doc.add_heading(_sanitize(section.heading), level=heading_level)

        # Content — render markdown to rich DOCX elements
        if section.content:
            _render_markdown_content(doc, section.content)

        # Recurse into subsections
        if section.subsections:
            _render_sections(doc, section.subsections)
