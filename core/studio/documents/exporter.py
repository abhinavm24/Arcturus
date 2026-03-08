import io
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from core.schemas.studio_schema import DocumentContentTree

def _add_section_to_doc(doc, section_data):
    """Recursively add document sections."""
    # Headings mapping (0 is typically Title, 1 is Heading 1, etc.)
    level = min(section_data.level, 9)
    doc.add_heading(section_data.heading, level=level)
    
    if section_data.content:
        # Check if content has newlines for multiple paragraphs
        for paragraph in str(section_data.content).split("\n"):
            p_text = paragraph.strip()
            if p_text:
                doc.add_paragraph(p_text)
                
    if section_data.citations:
        p = doc.add_paragraph()
        p.add_run(f"[Citations: {', '.join(section_data.citations)}]").italic = True

    for subsection in section_data.subsections:
        _add_section_to_doc(doc, subsection)


def export_docx(content_tree: DocumentContentTree) -> bytes:
    """Export a document content tree to DOCX format."""
    doc = Document()

    # Title
    doc.add_heading(content_tree.doc_title, 0)
    
    # Abstract
    if content_tree.abstract:
        doc.add_heading('Abstract', level=1)
        doc.add_paragraph(content_tree.abstract)
        
    # Main Sections
    for section in content_tree.sections:
        _add_section_to_doc(doc, section)
        
    # Bibliography
    if content_tree.bibliography:
        doc.add_page_break()
        doc.add_heading('Bibliography', level=1)
        for bib in content_tree.bibliography:
            key = bib.get("key", "")
            title = bib.get("title", "")
            author = bib.get("author", "")
            doc.add_paragraph(f"[{key}] {author}. {title}.")

    doc_stream = io.BytesIO()
    doc.save(doc_stream)
    doc_stream.seek(0)
    return doc_stream.read()
